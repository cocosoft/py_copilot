"""
文件读取工具

官方工具：提供本地文件读取功能，支持多种格式
"""

from typing import Dict, Any, List, Optional
import os
import time
import logging
from pathlib import Path

from app.modules.function_calling.base_tool import ServiceTool
from app.modules.function_calling.tool_schemas import (
    ToolParameter,
    ToolMetadata,
    ToolExecutionResult,
    ToolCategory
)

logger = logging.getLogger(__name__)


class FileReaderTool(ServiceTool):
    """
    文件读取工具
    
    官方工具：读取本地文件内容，支持多种格式（txt, md, pdf, docx, json, csv等）
    """
    
    # 支持的文本文件扩展名
    TEXT_EXTENSIONS = {'.txt', '.md', '.markdown', '.json', '.csv', '.xml', '.html', '.htm', '.yaml', '.yml', '.ini', '.conf', '.log', '.py', '.js', '.java', '.cpp', '.c', '.h', '.css', '.sql'}
    
    # 支持的二进制文件扩展名（需要特殊处理）
    BINARY_EXTENSIONS = {'.pdf', '.docx', '.doc', '.xlsx', '.xls', '.pptx', '.ppt'}
    
    def __init__(self):
        """初始化文件读取工具"""
        super().__init__()
    
    def _get_service(self):
        """
        获取服务实例
        
        Returns:
            None: 文件读取工具不需要外部服务
        """
        return None
    
    def _get_metadata(self) -> ToolMetadata:
        """
        获取工具元数据
        
        Returns:
            ToolMetadata: 工具元数据
        """
        return ToolMetadata(
            name="file_reader",
            display_name="文件读取工具",
            description="读取本地文件内容，支持多种格式（txt, md, pdf, docx, json, csv等）",
            category=ToolCategory.FILE.value,
            version="1.0.0",
            author="Py Copilot Team",
            icon="📄",
            tags=["文件", "读取", "官方"],
            is_active=True
        )
    
    def _get_parameters(self) -> List[ToolParameter]:
        """
        获取工具参数定义
        
        Returns:
            List[ToolParameter]: 参数定义列表
        """
        return [
            ToolParameter(
                name="file_path",
                type="string",
                description="要读取的文件路径（绝对路径或相对路径）",
                required=True
            ),
            ToolParameter(
                name="encoding",
                type="string",
                description="文件编码格式（默认自动检测）",
                required=False,
                default="auto"
            ),
            ToolParameter(
                name="max_size",
                type="integer",
                description="最大读取字节数（默认10MB）",
                required=False,
                default=10485760
            )
        ]
    
    def _detect_encoding(self, file_path: str) -> str:
        """
        检测文件编码
        
        Args:
            file_path: 文件路径
            
        Returns:
            str: 检测到的编码格式
        """
        try:
            # 尝试使用 chardet 检测编码
            import chardet
            with open(file_path, 'rb') as f:
                raw_data = f.read(4096)  # 读取前4KB用于检测
                result = chardet.detect(raw_data)
                encoding = result.get('encoding', 'utf-8')
                if encoding:
                    return encoding
        except ImportError:
            logger.debug("chardet 模块未安装，使用默认编码")
        except Exception as e:
            logger.debug(f"编码检测失败: {e}")
        
        # 默认使用 utf-8
        return 'utf-8'
    
    def _read_text_file(self, file_path: str, encoding: str, max_size: int) -> Dict[str, Any]:
        """
        读取文本文件
        
        Args:
            file_path: 文件路径
            encoding: 文件编码
            max_size: 最大读取字节数
            
        Returns:
            Dict[str, Any]: 读取结果
        """
        try:
            # 自动检测编码
            if encoding == 'auto':
                encoding = self._detect_encoding(file_path)
            
            # 读取文件内容
            with open(file_path, 'r', encoding=encoding, errors='replace') as f:
                content = f.read(max_size)
            
            # 检查是否被截断
            file_size = os.path.getsize(file_path)
            is_truncated = file_size > max_size
            
            return {
                'success': True,
                'content': content,
                'encoding': encoding,
                'file_size': file_size,
                'read_size': len(content.encode(encoding)),
                'is_truncated': is_truncated,
                'file_type': 'text'
            }
            
        except Exception as e:
            logger.error(f"读取文本文件失败: {e}")
            return {
                'success': False,
                'error': f"读取文件失败: {str(e)}"
            }
    
    def _read_pdf_file(self, file_path: str, max_size: int) -> Dict[str, Any]:
        """
        读取 PDF 文件
        
        Args:
            file_path: 文件路径
            max_size: 最大读取字节数
            
        Returns:
            Dict[str, Any]: 读取结果
        """
        try:
            # 尝试使用 PyPDF2 读取 PDF
            import PyPDF2
            
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                num_pages = len(pdf_reader.pages)
                
                content = ""
                total_chars = 0
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    # 检查是否超过最大大小
                    if total_chars + len(page_text) > max_size:
                        remaining = max_size - total_chars
                        content += page_text[:remaining]
                        break
                    
                    content += page_text + "\n"
                    total_chars += len(page_text) + 1
                
                return {
                    'success': True,
                    'content': content,
                    'file_type': 'pdf',
                    'num_pages': num_pages,
                    'is_truncated': total_chars >= max_size,
                    'file_size': os.path.getsize(file_path)
                }
                
        except ImportError:
            return {
                'success': False,
                'error': "读取 PDF 需要安装 PyPDF2 模块: pip install PyPDF2"
            }
        except Exception as e:
            logger.error(f"读取 PDF 文件失败: {e}")
            return {
                'success': False,
                'error': f"读取 PDF 失败: {str(e)}"
            }
    
    def _read_docx_file(self, file_path: str, max_size: int) -> Dict[str, Any]:
        """
        读取 Word 文档
        
        Args:
            file_path: 文件路径
            max_size: 最大读取字节数
            
        Returns:
            Dict[str, Any]: 读取结果
        """
        try:
            # 尝试使用 python-docx 读取 Word 文档
            import docx
            
            doc = docx.Document(file_path)
            content = []
            total_chars = 0
            
            for para in doc.paragraphs:
                text = para.text
                if total_chars + len(text) > max_size:
                    remaining = max_size - total_chars
                    content.append(text[:remaining])
                    break
                content.append(text)
                total_chars += len(text) + 1
            
            return {
                'success': True,
                'content': '\n'.join(content),
                'file_type': 'docx',
                'num_paragraphs': len(doc.paragraphs),
                'is_truncated': total_chars >= max_size,
                'file_size': os.path.getsize(file_path)
            }
            
        except ImportError:
            return {
                'success': False,
                'error': "读取 Word 文档需要安装 python-docx 模块: pip install python-docx"
            }
        except Exception as e:
            logger.error(f"读取 Word 文档失败: {e}")
            return {
                'success': False,
                'error': f"读取 Word 文档失败: {str(e)}"
            }
    
    async def execute(self, parameters: Dict[str, Any]) -> ToolExecutionResult:
        """
        执行文件读取操作
        
        Args:
            parameters: 执行参数
                - file_path: 文件路径
                - encoding: 编码格式（可选）
                - max_size: 最大读取字节数（可选）
        
        Returns:
            ToolExecutionResult: 执行结果
        """
        start_time = time.time()
        
        try:
            # 获取参数
            file_path = parameters.get('file_path')
            encoding = parameters.get('encoding', 'auto')
            max_size = parameters.get('max_size', 10485760)  # 默认10MB
            
            # 验证参数
            if not file_path:
                return ToolExecutionResult(
                    success=False,
                    error="缺少必需参数: file_path",
                    execution_time=time.time() - start_time,
                    tool_name="file_reader"
                )
            
            # 展开用户主目录
            file_path = os.path.expanduser(file_path)
            
            # 检查文件是否存在
            if not os.path.exists(file_path):
                return ToolExecutionResult(
                    success=False,
                    error=f"文件不存在: {file_path}",
                    execution_time=time.time() - start_time,
                    tool_name="file_reader"
                )
            
            # 检查是否是文件
            if not os.path.isfile(file_path):
                return ToolExecutionResult(
                    success=False,
                    error=f"路径不是文件: {file_path}",
                    execution_time=time.time() - start_time,
                    tool_name="file_reader"
                )
            
            # 获取文件扩展名
            _, ext = os.path.splitext(file_path)
            ext = ext.lower()
            
            # 根据文件类型选择读取方式
            if ext in self.TEXT_EXTENSIONS or ext == '':
                result = self._read_text_file(file_path, encoding, max_size)
            elif ext == '.pdf':
                result = self._read_pdf_file(file_path, max_size)
            elif ext in {'.docx', '.doc'}:
                result = self._read_docx_file(file_path, max_size)
            else:
                # 默认按文本文件读取
                result = self._read_text_file(file_path, encoding, max_size)
            
            execution_time = time.time() - start_time
            
            if result.get('success'):
                return ToolExecutionResult(
                    success=True,
                    result={
                        'content': result.get('content'),
                        'file_type': result.get('file_type', 'unknown'),
                        'file_size': result.get('file_size'),
                        'encoding': result.get('encoding'),
                        'is_truncated': result.get('is_truncated', False)
                    },
                    execution_time=execution_time,
                    tool_name="file_reader",
                    metadata={
                        'file_path': file_path,
                        'original_size': result.get('file_size'),
                        'read_size': result.get('read_size', len(result.get('content', '')))
                    }
                )
            else:
                return ToolExecutionResult(
                    success=False,
                    error=result.get('error', '未知错误'),
                    execution_time=execution_time,
                    tool_name="file_reader"
                )

        except Exception as e:
            logger.error(f"文件读取工具执行失败: {e}")
            return ToolExecutionResult(
                success=False,
                error=f"执行失败: {str(e)}",
                execution_time=time.time() - start_time,
                tool_name="file_reader"
            )
